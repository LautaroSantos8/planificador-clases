"""
Exportador de Planificaciones
Genera archivos PDF y DOCX a partir de las respuestas del asistente IA.

Convierte el Markdown generado por Gemini a documentos bien formateados
con el encabezado de la Escuela Municipal Dr. Jorge Orgaz.
"""

import re
import io
from datetime import datetime

# ─── PDF ────────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable, KeepTogether
)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY

# ─── DOCX ───────────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================================
# COLORES DE LA INSTITUCIÓN
# ============================================================================
COLOR_PRIMARIO_HEX = "#3730A3"   # Indigo-800 — coincide con el front
COLOR_SECUNDARIO_HEX = "#6366F1" # Indigo-500
COLOR_NIVEL_NEE = "#DC2626"      # Rojo
COLOR_NIVEL_LP  = "#D97706"      # Amarillo-ocre
COLOR_NIVEL_LE  = "#16A34A"      # Verde

def _hex_to_rgb(hex_color: str):
    """Convierte #RRGGBB a (R, G, B) en 0-255."""
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


# ============================================================================
# PARSER DE MARKDOWN SIMPLE
# ============================================================================
def _parse_markdown(text: str) -> list:
    """
    Parsea el markdown del asistente y devuelve una lista de bloques:
    [("heading1", texto), ("heading2", texto), ("heading3", texto),
     ("bold", texto), ("normal", texto), ("bullet", texto),
     ("hr", None), ("table_header", [cols]), ("table_row", [cols])]
    """
    blocks = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Salto de línea vacío
        if not line.strip():
            i += 1
            continue

        # Separador ---
        if re.match(r"^-{3,}$", line.strip()):
            blocks.append(("hr", None))
            i += 1
            continue

        # Encabezados
        m = re.match(r"^(#{1,3})\s+(.+)", line)
        if m:
            level = len(m.group(1))
            blocks.append((f"heading{level}", m.group(2).strip()))
            i += 1
            continue

        # Tablas markdown: | col | col |
        if line.strip().startswith("|"):
            # Recolectar todas las filas de la tabla
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                raw = lines[i].strip()
                # Ignorar líneas separadoras |---|---|
                if re.match(r"^\|[\s\-\|:]+\|$", raw):
                    i += 1
                    continue
                cols = [c.strip() for c in raw.strip("|").split("|")]
                table_lines.append(cols)
                i += 1
            if table_lines:
                blocks.append(("table_header", table_lines[0]))
                for row in table_lines[1:]:
                    blocks.append(("table_row", row))
            continue

        # Bullet: - o *
        m = re.match(r"^[\-\*]\s+(.+)", line)
        if m:
            blocks.append(("bullet", m.group(1).strip()))
            i += 1
            continue

        # Bold solo: **texto**
        if re.match(r"^\*\*[^*]+\*\*$", line.strip()):
            text_clean = line.strip().strip("*")
            blocks.append(("bold", text_clean))
            i += 1
            continue

        # Párrafo normal
        blocks.append(("normal", line.strip()))
        i += 1

    return blocks


def _strip_md_inline(text: str) -> str:
    """Elimina marcadores inline de markdown (bold, italic, code)."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _md_inline_to_reportlab(text: str) -> str:
    """Convierte markdown inline a tags de ReportLab."""
    # Bold: **texto** → <b>texto</b>
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    # Italic: *texto* → <i>texto</i>
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    # Code: `texto` → <font face="Courier">texto</font>
    text = re.sub(r"`(.+?)`", r'<font face="Courier">\1</font>', text)
    # Escapar & que no son entidades
    text = re.sub(r"&(?![a-zA-Z]+;|#\d+;)", "&amp;", text)
    return text


# ============================================================================
# GENERADOR PDF
# ============================================================================
def generar_pdf(
    respuesta_md: str,
    nombre_docente: str,
    grado: str,
    division: str,
    materia: str,
    consulta_original: str,
    fecha: datetime = None,
) -> bytes:
    """
    Genera un PDF bien formateado a partir del markdown del asistente.
    Retorna los bytes del PDF.
    """
    fecha = fecha or datetime.now()
    buffer = io.BytesIO()

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
        title=f"Planificación {grado}° {division} - {materia}",
        author=nombre_docente,
    )

    # ── Estilos ──────────────────────────────────────────────────────────────
    styles = getSampleStyleSheet()
    R, G, B = _hex_to_rgb(COLOR_PRIMARIO_HEX)
    primario = colors.Color(R/255, G/255, B/255)
    R2, G2, B2 = _hex_to_rgb(COLOR_SECUNDARIO_HEX)
    secundario = colors.Color(R2/255, G2/255, B2/255)

    st_titulo_doc = ParagraphStyle(
        "TituloDoc", parent=styles["Title"],
        fontSize=16, textColor=primario,
        spaceAfter=4, alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )
    st_subtitulo = ParagraphStyle(
        "Subtitulo", parent=styles["Normal"],
        fontSize=10, textColor=colors.HexColor("#4B5563"),
        spaceAfter=2, alignment=TA_CENTER,
    )
    st_h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"],
        fontSize=14, textColor=primario,
        spaceBefore=14, spaceAfter=6,
        fontName="Helvetica-Bold",
        borderPad=4,
    )
    st_h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"],
        fontSize=12, textColor=secundario,
        spaceBefore=10, spaceAfter=4,
        fontName="Helvetica-Bold",
    )
    st_h3 = ParagraphStyle(
        "H3", parent=styles["Heading3"],
        fontSize=11, textColor=colors.HexColor("#374151"),
        spaceBefore=8, spaceAfter=3,
        fontName="Helvetica-Bold",
    )
    st_normal = ParagraphStyle(
        "Normal2", parent=styles["Normal"],
        fontSize=10, leading=14,
        spaceAfter=4, alignment=TA_JUSTIFY,
    )
    st_bold = ParagraphStyle(
        "Bold", parent=styles["Normal"],
        fontSize=10, leading=14,
        spaceAfter=4, fontName="Helvetica-Bold",
    )
    st_bullet = ParagraphStyle(
        "Bullet", parent=styles["Normal"],
        fontSize=10, leading=14,
        leftIndent=16, spaceAfter=2,
        bulletIndent=4,
    )
    st_consulta = ParagraphStyle(
        "Consulta", parent=styles["Normal"],
        fontSize=9, leading=13,
        textColor=colors.HexColor("#6B7280"),
        leftIndent=8, rightIndent=8,
        borderPad=6, spaceAfter=8,
        backColor=colors.HexColor("#F3F4F6"),
    )

    story = []

    # ── Encabezado ───────────────────────────────────────────────────────────
    story.append(Paragraph("Escuela Municipal Dr. Jorge Orgaz", st_titulo_doc))
    story.append(Paragraph("Villa Rivera Indarte, Córdoba · Planificador Docente IA", st_subtitulo))
    story.append(HRFlowable(width="100%", thickness=2, color=primario, spaceAfter=8))

    # Metadatos en tabla
    meta_data = [
        ["Docente:", nombre_docente, "Fecha:", fecha.strftime("%d/%m/%Y")],
        ["Grado/Div:", f"{grado}° {division}", "Materia:", materia],
    ]
    meta_table = Table(meta_data, colWidths=[3*cm, 6.5*cm, 2.5*cm, 5.5*cm])
    meta_table.setStyle(TableStyle([
        ("FONTNAME", (0,0), (-1,-1), "Helvetica"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("FONTNAME", (0,0), (0,-1), "Helvetica-Bold"),
        ("FONTNAME", (2,0), (2,-1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0), (-1,-1), colors.HexColor("#374151")),
        ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ("TOPPADDING", (0,0), (-1,-1), 3),
    ]))
    story.append(meta_table)
    story.append(Spacer(1, 8))

    # Consulta original
    story.append(Paragraph(f"<b>Consulta:</b> {_md_inline_to_reportlab(consulta_original)}", st_consulta))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E5E7EB"), spaceAfter=10))

    # ── Contenido desde Markdown ─────────────────────────────────────────────
    blocks = _parse_markdown(respuesta_md)

    # Acumular filas de tabla en curso
    current_table_headers = None
    current_table_rows = []

    def flush_table():
        nonlocal current_table_headers, current_table_rows
        if current_table_headers is None:
            return
        all_rows = [current_table_headers] + current_table_rows
        # Igualar columnas
        n_cols = max(len(r) for r in all_rows)
        for r in all_rows:
            while len(r) < n_cols:
                r.append("")
        # Ancho dinámico
        avail = 17 * cm
        col_w = avail / n_cols
        col_widths = [col_w] * n_cols

        t_data = []
        for row in all_rows:
            t_data.append([Paragraph(_md_inline_to_reportlab(c), ParagraphStyle(
                "tc", parent=styles["Normal"],
                fontSize=8, leading=11,
            )) for c in row])

        t = Table(t_data, colWidths=col_widths, repeatRows=1)
        R_p, G_p, B_p = _hex_to_rgb(COLOR_PRIMARIO_HEX)
        header_color = colors.Color(R_p/255, G_p/255, B_p/255)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), header_color),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white),
            ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("FONTSIZE", (0,0), (-1,0), 8),
            ("ALIGN", (0,0), (-1,-1), "LEFT"),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [colors.white, colors.HexColor("#F9FAFB")]),
            ("GRID", (0,0), (-1,-1), 0.5, colors.HexColor("#D1D5DB")),
            ("BOTTOMPADDING", (0,0), (-1,-1), 5),
            ("TOPPADDING", (0,0), (-1,-1), 5),
            ("LEFTPADDING", (0,0), (-1,-1), 6),
            ("RIGHTPADDING", (0,0), (-1,-1), 6),
        ]))
        story.append(t)
        story.append(Spacer(1, 6))
        current_table_headers = None
        current_table_rows = []

    for block_type, content in blocks:
        # Si viene algo que no es fila de tabla, volcar tabla pendiente
        if block_type not in ("table_header", "table_row"):
            flush_table()

        if block_type == "heading1":
            story.append(Paragraph(_md_inline_to_reportlab(content), st_h1))
        elif block_type == "heading2":
            story.append(Paragraph(_md_inline_to_reportlab(content), st_h2))
        elif block_type == "heading3":
            story.append(Paragraph(_md_inline_to_reportlab(content), st_h3))
        elif block_type == "bold":
            story.append(Paragraph(_md_inline_to_reportlab(content), st_bold))
        elif block_type == "normal":
            story.append(Paragraph(_md_inline_to_reportlab(content), st_normal))
        elif block_type == "bullet":
            story.append(Paragraph(
                f"• {_md_inline_to_reportlab(content)}", st_bullet
            ))
        elif block_type == "hr":
            story.append(HRFlowable(
                width="100%", thickness=1,
                color=colors.HexColor("#E5E7EB"),
                spaceBefore=6, spaceAfter=6,
            ))
        elif block_type == "table_header":
            current_table_headers = list(content)
            current_table_rows = []
        elif block_type == "table_row":
            current_table_rows.append(list(content))

    flush_table()

    # ── Cuadro de observaciones ──────────────────────────────────────────────
    story.append(Spacer(1, 16))
    story.append(Paragraph(
        "Observaciones del docente",
        ParagraphStyle("obs_titulo", parent=styles["Normal"],
                    fontSize=10, fontName="Helvetica-Bold",
                    textColor=primario, spaceAfter=4)
    ))
    obs_data = [["Anotá aquí tus observaciones sobre cómo resultó la clase:"],
                [" "], [" "], [" "], [" "]]
    obs_table = Table(obs_data, colWidths=[17*cm])
    obs_table.setStyle(TableStyle([
        ("BOX", (0,0), (-1,-1), 1, colors.HexColor("#D1D5DB")),
        ("INNERGRID", (0,0), (-1,-1), 0.5, colors.HexColor("#E5E7EB")),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("TEXTCOLOR", (0,0), (0,0), colors.HexColor("#6B7280")),
        ("TOPPADDING", (0,0), (-1,-1), 8),
        ("BOTTOMPADDING", (0,0), (-1,-1), 8),
        ("LEFTPADDING", (0,0), (-1,-1), 8),
        ("MINROWHEIGHT", (0,1), (-1,-1), 20),
    ]))
    story.append(obs_table)

    # ── Pie de página ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 12))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#E5E7EB")))
    story.append(Paragraph(
        f"Generado por el Asistente IA · {fecha.strftime('%d/%m/%Y %H:%M')}",
        ParagraphStyle("pie", parent=styles["Normal"],
                       fontSize=8, textColor=colors.HexColor("#9CA3AF"),
                       alignment=TA_CENTER, spaceBefore=4)
    ))

    doc.build(story)
    return buffer.getvalue()


# ============================================================================
# GENERADOR DOCX
# ============================================================================
def _set_cell_bg(cell, hex_color: str):
    """Pone color de fondo a una celda de tabla DOCX."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _add_run_bold(paragraph, text: str):
    run = paragraph.add_run(text)
    run.bold = True
    return run


def _md_inline_to_docx(paragraph, text: str):
    """
    Agrega runs a un párrafo DOCX interpretando **bold** y texto normal.
    """
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Limpiar otros marcadores
            clean = re.sub(r"\*(.+?)\*", r"\1", part)
            clean = re.sub(r"`(.+?)`", r"\1", clean)
            paragraph.add_run(clean)


def generar_docx(
    respuesta_md: str,
    nombre_docente: str,
    grado: str,
    division: str,
    materia: str,
    consulta_original: str,
    fecha: datetime = None,
) -> bytes:
    """
    Genera un archivo DOCX bien formateado a partir del markdown del asistente.
    Retorna los bytes del DOCX.
    """
    fecha = fecha or datetime.now()
    doc = Document()

    # ── Márgenes ──────────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    prim_r, prim_g, prim_b = _hex_to_rgb(COLOR_PRIMARIO_HEX)
    sec_r,  sec_g,  sec_b  = _hex_to_rgb(COLOR_SECUNDARIO_HEX)
    primario  = RGBColor(prim_r, prim_g, prim_b)
    secundario = RGBColor(sec_r, sec_g, sec_b)

    # ── Encabezado ───────────────────────────────────────────────────────────
    h = doc.add_heading("Escuela Municipal Dr. Jorge Orgaz", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = primario
        run.font.size = Pt(16)

    sub = doc.add_paragraph("Villa Rivera Indarte, Córdoba · Planificador Docente IA")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    # Línea divisoria via borde inferior
    p_border = doc.add_paragraph()
    p_border_fmt = p_border.paragraph_format
    p_border.paragraph_format.space_after = Pt(6)
    pPr = p_border._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), COLOR_PRIMARIO_HEX.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)

    # Tabla de metadatos
    meta_tbl = doc.add_table(rows=2, cols=4)
    meta_tbl.style = "Table Grid"
    meta_data = [
        ["Docente:", nombre_docente, "Fecha:", fecha.strftime("%d/%m/%Y")],
        ["Grado/Div:", f"{grado}° {division}", "Materia:", materia],
    ]
    for r_idx, row_data in enumerate(meta_data):
        row = meta_tbl.rows[r_idx]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p = cell.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            if c_idx % 2 == 0:
                run.bold = True
                run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    doc.add_paragraph()  # Espacio

    # Consulta original
    p_consulta = doc.add_paragraph()
    p_consulta.paragraph_format.left_indent = Cm(0.5)
    run_label = p_consulta.add_run("Consulta: ")
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_consulta = p_consulta.add_run(_strip_md_inline(consulta_original))
    run_consulta.font.size = Pt(9)
    run_consulta.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.add_paragraph()

    # ── Contenido desde Markdown ─────────────────────────────────────────────
    blocks = _parse_markdown(respuesta_md)

    current_table_headers = None
    current_table_rows = []

    def flush_table_docx():
        nonlocal current_table_headers, current_table_rows
        if current_table_headers is None:
            return
        all_rows = [current_table_headers] + current_table_rows
        n_cols = max(len(r) for r in all_rows)
        for r in all_rows:
            while len(r) < n_cols:
                r.append("")

        tbl = doc.add_table(rows=len(all_rows), cols=n_cols)
        tbl.style = "Table Grid"

        for r_idx, row_data in enumerate(all_rows):
            row = tbl.rows[r_idx]
            for c_idx, val in enumerate(row_data):
                cell = row.cells[c_idx]
                p = cell.paragraphs[0]
                run = p.add_run(_strip_md_inline(val))
                run.font.size = Pt(8)
                if r_idx == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    _set_cell_bg(cell, COLOR_PRIMARIO_HEX)

        doc.add_paragraph()
        current_table_headers = None
        current_table_rows.clear()

    for block_type, content in blocks:
        if block_type not in ("table_header", "table_row"):
            flush_table_docx()

        if block_type == "heading1":
            h = doc.add_heading(level=1)
            h.clear()
            run = h.add_run(_strip_md_inline(content))
            run.font.color.rgb = primario
            run.font.size = Pt(14)
            run.bold = True

        elif block_type == "heading2":
            h = doc.add_heading(level=2)
            h.clear()
            run = h.add_run(_strip_md_inline(content))
            run.font.color.rgb = secundario
            run.font.size = Pt(12)
            run.bold = True

        elif block_type == "heading3":
            h = doc.add_heading(level=3)
            h.clear()
            run = h.add_run(_strip_md_inline(content))
            run.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
            run.font.size = Pt(11)
            run.bold = True

        elif block_type == "bold":
            p = doc.add_paragraph()
            run = p.add_run(_strip_md_inline(content))
            run.bold = True
            run.font.size = Pt(10)

        elif block_type == "normal":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _md_inline_to_docx(p, content)
            for run in p.runs:
                run.font.size = Pt(10)

        elif block_type == "bullet":
            p = doc.add_paragraph(style="List Bullet")
            _md_inline_to_docx(p, content)
            for run in p.runs:
                run.font.size = Pt(10)

        elif block_type == "hr":
            p = doc.add_paragraph()
            pPr2 = p._p.get_or_add_pPr()
            pBdr2 = OxmlElement("w:pBdr")
            bot2 = OxmlElement("w:bottom")
            bot2.set(qn("w:val"), "single")
            bot2.set(qn("w:sz"), "4")
            bot2.set(qn("w:space"), "1")
            bot2.set(qn("w:color"), "E5E7EB")
            pBdr2.append(bot2)
            pPr2.append(pBdr2)

        elif block_type == "table_header":
            current_table_headers = list(content)
            current_table_rows = []

        elif block_type == "table_row":
            current_table_rows.append(list(content))

    flush_table_docx()

    # ── Cuadro de observaciones ──────────────────────────────────────────────
    doc.add_paragraph()
    p_obs_titulo = doc.add_paragraph()
    run_obs = p_obs_titulo.add_run("Observaciones del docente")
    run_obs.bold = True
    run_obs.font.size = Pt(10)
    run_obs.font.color.rgb = primario

    obs_tbl = doc.add_table(rows=6, cols=1)
    obs_tbl.style = "Table Grid"
    # Primera fila con indicación
    obs_tbl.rows[0].cells[0].paragraphs[0].add_run(
        "Anotá aquí tus observaciones sobre cómo resultó la clase:"
    ).font.size = Pt(9)
    obs_tbl.rows[0].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
    # Filas vacías para escribir
    for i in range(1, 6):
        cell = obs_tbl.rows[i].cells[0]
        cell.paragraphs[0].add_run(" ")
        from docx.shared import Cm as DocxCm
        cell._tc.get_or_add_tcPr()
    doc.add_paragraph()

    # ── Pie ───────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    pie = doc.add_paragraph(
        f"Generado por el Asistente IA · {fecha.strftime('%d/%m/%Y %H:%M')}"
    )
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.runs[0].font.size = Pt(8)
    pie.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
