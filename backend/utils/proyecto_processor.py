"""
Procesador de Proyectos Docentes - Versión Simplificada
========================================================

Procesa proyectos áulicos/institucionales y planificaciones anuales
en cualquier formato (.docx, .pdf, .xlsx) y los divide en chunks 
por tamaño para almacenar en ChromaDB.

No asume estructura específica - funciona con cualquier formato.
"""

import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Optional, Union

# Dependencias opcionales
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from PyPDF2 import PdfReader
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    import openpyxl
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False
    
import logging
logger = logging.getLogger(__name__)

@dataclass
class ProyectoMetadata:
    """Metadata de un chunk de proyecto."""
    proyecto_id: str
    titulo: str
    docente_id: int
    institucion_id: int
    grado: str
    tipo: str = "proyecto_aulico"  # proyecto_aulico | planificacion_anual
    materias: List[str] = field(default_factory=list)
    año_proyecto: int = 2025
    chunk_index: int = 0
    total_chunks: int = 1


@dataclass
class ProyectoChunk:
    """Chunk de texto con su metadata."""
    texto: str
    metadata: ProyectoMetadata


# Patrones para detectar materias — requieren menciones múltiples o en contexto clave
PATRONES_MATERIA = {
    'matematicas': [r'matem[aá]tica', r'c[aá]lculo', r'geometr[ií]a'],
    'lengua': [r'lengua', r'literatura', r'lectura y escritura', r'pr[aá]cticas del lenguaje'],
    'ciencias_naturales': [r'ciencias?\s*naturales?', r'biolog[ií]a', r'medio\s*ambiente'],
    'ciencias_sociales': [r'ciencias?\s*sociales?', r'historia', r'geograf[ií]a'],
    'educacion_fisica': [r'educaci[oó]n\s*f[ií]sica'],
    'educacion_artistica': [r'educaci[oó]n\s*art[ií]stica', r'expresi[oó]n\s*art[ií]stica', r'artes?\s*visuales?', r'm[uú]sica'],
    'tecnologia': [r'\btecnolog[ií]a\s+(?!productos)', r'inform[aá]tica', r'\btics\b'],
    'ciudadania': [r'ciudadan[ií]a', r'formaci[oó]n\s*[eé]tica', r'educaci[oó]n\s*sexual\s*integral', r'\besi\b'],
}

# Patrones de limpieza
PATRONES_RUIDO = [
    r'!\[.*?\]\(media/.*?\)',  # Imágenes en markdown
    r'\{width=.*?\}',
    r'\{height=.*?\}',
]


class ProyectoProcessor:
    """
    Procesador simplificado de proyectos docentes.
    
    Divide el documento en chunks por tamaño, sin asumir estructura específica.
    """
    
    def __init__(self, chunk_size: int = 1500, chunk_overlap: int = 200):
        """
        Args:
            chunk_size: Tamaño aproximado de cada chunk en caracteres
            chunk_overlap: Solapamiento entre chunks para mantener contexto
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.stats = {
            'formato': None,
            'caracteres_totales': 0,
            'chunks_generados': 0,
        }
    
    def _extract_from_docx(self, file_path: Union[str, Path]) -> str:
        """Extrae texto de archivo Word."""
        if not DOCX_SUPPORT:
            raise ImportError("Instalar python-docx: pip install python-docx")
        
        doc = Document(str(file_path))
        textos = []
        
        # Extraer párrafos
        for para in doc.paragraphs:
            texto = para.text.strip()
            if texto:
                textos.append(texto)
        
        # Extraer tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    texto = cell.text.strip()
                    if texto:
                        row_text.append(texto)
                if row_text:
                    textos.append(' | '.join(row_text))
        
        return '\n'.join(textos)
    
    def _extract_from_pdf(self, file_path: Union[str, Path]) -> str:
        """Extrae texto de archivo PDF."""
        if not PDF_SUPPORT:
            raise ImportError("Instalar PyPDF2: pip install PyPDF2")
        
        reader = PdfReader(str(file_path))
        textos = []
        
        for page in reader.pages:
            texto = page.extract_text()
            if texto:
                textos.append(texto.strip())
        
        return '\n'.join(textos)
    
    def _extract_from_xlsx(self, file_path: Union[str, Path]) -> str:
        """Extrae texto de archivo Excel."""
        if not XLSX_SUPPORT:
            raise ImportError("Instalar openpyxl: pip install openpyxl")
        
        wb = openpyxl.load_workbook(str(file_path), data_only=True)
        textos = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            textos.append(f"=== {sheet_name} ===")
            
            for row in sheet.iter_rows():
                row_values = []
                for cell in row:
                    if cell.value is not None:
                        valor = str(cell.value).strip()
                        if valor:
                            row_values.append(valor)
                
                if row_values:
                    textos.append(' | '.join(row_values))
        
        wb.close()
        return '\n'.join(textos)
    
    def _extract_planificacion_anual(self, file_path) -> str:
        """
        Extrae y estructura una planificación anual usando Gemini.
        Funciona con cualquier formato de tabla que use el docente.
        """
        from google import genai
        import os
        import time

        extension = Path(file_path).suffix.lower()
        
        # Para docx, extraer tablas de forma estructurada
        if extension == '.docx':
            from docx import Document
            doc = Document(str(file_path))
            
            partes = []
            # Texto de párrafos
            for p in doc.paragraphs:
                if p.text.strip():
                    partes.append(p.text.strip())
            
            # Tablas con formato estructurado
            for table in doc.tables:
                if len(table.rows) < 2:
                    continue
                # Obtener headers de la primera fila
                headers = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
                partes.append(f"\nTABLA CON COLUMNAS: {' | '.join(headers)}")
                
                for row in table.rows[1:]:
                    cells = [c.text.strip().replace('\n', ' ')[:200] for c in row.cells]
                    fila_formateada = []
                    for h, c in zip(headers, cells):
                        if c:
                            fila_formateada.append(f"{h}={c}")
                    if fila_formateada:
                        partes.append("FILA: " + " | ".join(fila_formateada))
            
            texto_crudo = '\n'.join(partes)
        elif extension == '.pdf':
            texto_crudo = self._extract_from_pdf(file_path)
        else:
            texto_crudo = self._extract_from_xlsx(file_path)

        if not texto_crudo.strip():
            return texto_crudo

        # Usar Gemini para estructurar el contenido
        try:
            client = genai.Client(api_key=os.environ.get('GEMINI_API_KEY'))

            prompt = f"""Este documento es una planificación anual docente argentina de nivel primario.
Extraé toda la información y devolvela ÚNICAMENTE en este formato, un bloque por cada combinación de período+materia+eje:

PERÍODO: [mes o rango de meses, ej: ABRIL, MAYO-JUNIO, JULIO-AGOSTO]
MATERIA: [nombre de la materia, ej: MATEMÁTICA, LENGUA, CIENCIAS NATURALES]
EJE: [nombre del eje o unidad temática]
CONTENIDOS:
[contenidos e indicadores de logro completos]

---

Reglas importantes:
- Cada FILA de la tabla es un bloque separado con su propio período, materia y eje
- El campo TIEMPO de cada fila indica el PERÍODO
- El campo ESPACIO CURRICULAR indica la MATERIA
- Si un período abarca varios meses (ej: ABRIL – MAYO - JUNIO), escribilo tal cual
- Mantené los contenidos COMPLETOS, no los resumás ni acortés
- Separá cada bloque con ---
- No agregues explicaciones, introducciones ni texto extra — solo los bloques en el formato indicado

DOCUMENTO A PROCESAR:
{texto_crudo[:30000]}
"""

            for intento in range(3):
                try:
                    response = client.models.generate_content(
                        model='gemini-2.5-flash',
                        contents=prompt,
                    )
                    return response.text
                except Exception as e:
                    if '429' in str(e) and intento < 2:
                        espera = 5 * (2 ** intento)
                        logger.warning(f"Rate limit Gemini (intento {intento + 1}). Reintentando en {espera}s...")
                        time.sleep(espera)
                    else:
                        logger.warning(f"Gemini no pudo estructurar la planificación: {e}. Usando texto crudo.")
                        return texto_crudo

        except Exception as e:
            logger.warning(f"Error general procesando planificación: {e}. Usando texto crudo.")
            return texto_crudo

    def _limpiar_texto(self, texto: str) -> str:
        """Limpia el texto de ruido."""
        for patron in PATRONES_RUIDO:
            texto = re.sub(patron, '', texto)
        
        # Normalizar espacios y saltos de línea
        texto = re.sub(r'\n{3,}', '\n\n', texto)
        texto = re.sub(r' {2,}', ' ', texto)
        
        return texto.strip()
    
    def _detectar_materias(self, texto: str) -> List[str]:
        """
        Detecta materias con umbral mínimo de menciones para evitar falsos positivos.
        Una materia se detecta solo si aparece 2+ veces o en frases clave.
        """
        materias = set()
        texto_lower = texto.lower()

        # Frases que indican que el documento ES de esa materia (1 mención alcanza)
        FRASES_CLAVE = {
            'matematicas': [r'proyecto.{0,30}matem[aá]tica', r'matem[aá]tica.{0,30}grado', r'[aá]rea.{0,20}matem[aá]tica'],
            'lengua': [r'proyecto.{0,30}lengua', r'lengua.{0,30}grado', r'[aá]rea.{0,20}lengua', r'lectura y escritura'],
            'ciencias_naturales': [r'proyecto.{0,30}ciencias?\s*naturales?', r'ciencias?\s*naturales?.{0,30}grado'],
            'ciencias_sociales': [r'proyecto.{0,30}ciencias?\s*sociales?', r'ciencias?\s*sociales?.{0,30}grado'],
            'educacion_fisica': [r'proyecto.{0,30}educaci[oó]n\s*f[ií]sica'],
            'educacion_artistica': [r'proyecto.{0,30}educaci[oó]n\s*art[ií]stica', r'proyecto.{0,30}artes?\s*visuales?', r'proyecto.{0,30}m[uú]sica'],
            'tecnologia': [r'proyecto.{0,30}tecnolog[ií]a', r'[aá]rea.{0,20}tecnolog[ií]a'],
            'ciudadania': [r'proyecto.{0,30}esi', r'proyecto.{0,30}ciudadan[ií]a', r'educaci[oó]n\s*sexual\s*integral'],
        }

        for materia, patrones in PATRONES_MATERIA.items():
            # Verificar frases clave primero (1 mención alcanza)
            for frase in FRASES_CLAVE.get(materia, []):
                if re.search(frase, texto_lower):
                    materias.add(materia)
                    break

            if materia in materias:
                continue

            # Si no hay frase clave, requerir 2+ menciones
            menciones = 0
            for patron in patrones:
                menciones += len(re.findall(patron, texto_lower))
            if menciones >= 2:
                materias.add(materia)

        return sorted(list(materias))
    
    def _extraer_titulo(self, texto: str) -> str:
        """Extrae el título del proyecto de las primeras líneas."""
        lineas = texto.split('\n')
        titulo_partes = []
        
        for linea in lineas[:15]:
            linea = linea.strip()
            if not linea:
                continue
            
            # Ignorar líneas institucionales
            if re.match(r'^(escuela|colegio|instituto|docente|año:|grado:|área)', linea.lower()):
                continue
            
            # Líneas cortas al inicio suelen ser títulos
            if len(linea) > 5 and len(linea) < 150:
                titulo_partes.append(linea)
                if len(titulo_partes) >= 2:
                    break
        
        if titulo_partes:
            return ' - '.join(titulo_partes[:2])
        
        return "Proyecto sin título"
    
    def _generar_proyecto_id(self, titulo: str, docente_id: int) -> str:
        """Genera un ID único para el proyecto."""
        # Limpiar título para usar como ID
        titulo_limpio = re.sub(r'[^a-zA-Z0-9\s]', '', titulo.lower())
        titulo_limpio = '_'.join(titulo_limpio.split()[:4])
        return f"{titulo_limpio}_doc{docente_id}"
    
    def _dividir_en_chunks(self, texto: str) -> List[str]:
        """
        Divide el texto en chunks por tamaño.
        
        Intenta cortar en párrafos o puntos para mantener coherencia.
        """
        chunks = []
        
        # Dividir por párrafos primero
        parrafos = texto.split('\n\n')
        
        chunk_actual = ""
        
        for parrafo in parrafos:
            parrafo = parrafo.strip()
            if not parrafo:
                continue
            
            # Si agregar este párrafo excede el tamaño, guardar chunk actual
            if len(chunk_actual) + len(parrafo) > self.chunk_size and chunk_actual:
                chunks.append(chunk_actual.strip())
                # Mantener overlap del final del chunk anterior
                if len(chunk_actual) > self.chunk_overlap:
                    chunk_actual = chunk_actual[-self.chunk_overlap:]
                else:
                    chunk_actual = ""
            
            chunk_actual += "\n\n" + parrafo if chunk_actual else parrafo
        
        # Agregar último chunk si tiene contenido
        if chunk_actual.strip():
            chunks.append(chunk_actual.strip())
        
        # Si un chunk es muy largo, dividirlo por oraciones
        chunks_finales = []
        for chunk in chunks:
            if len(chunk) > self.chunk_size * 1.5:
                # Dividir por oraciones
                oraciones = re.split(r'(?<=[.!?])\s+', chunk)
                sub_chunk = ""
                for oracion in oraciones:
                    if len(sub_chunk) + len(oracion) > self.chunk_size and sub_chunk:
                        chunks_finales.append(sub_chunk.strip())
                        sub_chunk = ""
                    sub_chunk += " " + oracion if sub_chunk else oracion
                if sub_chunk.strip():
                    chunks_finales.append(sub_chunk.strip())
            else:
                chunks_finales.append(chunk)
        
        return chunks_finales
    
    def process_proyecto(
        self,
        file_path: Union[str, Path],
        docente_id: int,
        institucion_id: int,
        grado: str = None,
        materias: List[str] = None,
        año_proyecto: int = 2025,
        tipo: str = "proyecto_aulico",
    ) -> List[ProyectoChunk]:
        """
        Procesa un proyecto o planificación y genera chunks para ChromaDB.
        
        Args:
            file_path: Ruta al archivo .docx, .pdf o .xlsx
            docente_id: ID del docente
            institucion_id: ID de la institución
            grado: Grado (ej: "4", "5-6", "sala_5"). Si no se da, se detecta.
            materias: Lista de materias. Si no se da, se detectan.
            año_proyecto: Año del proyecto
            tipo: "proyecto_aulico" o "planificacion_anual"
        
        Returns:
            Lista de ProyectoChunk listos para ChromaDB
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        # Extraer texto según formato
        if extension == '.docx':
            self.stats['formato'] = 'docx'
            if tipo == 'planificacion_anual':
                texto = self._extract_planificacion_anual(file_path)
            else:
                texto = self._extract_from_docx(file_path)
        elif extension == '.pdf':
            self.stats['formato'] = 'pdf'
            texto = self._extract_from_pdf(file_path)
        elif extension in ['.xlsx', '.xls']:
            self.stats['formato'] = 'xlsx'
            texto = self._extract_from_xlsx(file_path)
        else:
            raise ValueError(f"Formato no soportado: {extension}. Use .docx, .pdf o .xlsx")
        
        # Limpiar texto
        texto = self._limpiar_texto(texto)
        self.stats['caracteres_totales'] = len(texto)
        
        # Extraer título
        titulo = self._extraer_titulo(texto)
        
        # Detectar materias si no se proporcionan
        if not materias:
            materias = self._detectar_materias(texto)
        
        # Generar ID del proyecto
        proyecto_id = self._generar_proyecto_id(titulo, docente_id)
        
        # Dividir en chunks
        if tipo == 'planificacion_anual' and '---' in texto:
            # Para planificación anual, cada bloque separado por --- es un chunk
            bloques = [b.strip() for b in texto.split('---') if b.strip()]
            chunks_texto = []
            for bloque in bloques:
                if len(bloque) > self.chunk_size * 2:
                    # Si un bloque es muy largo, subdividirlo manteniendo el header
                    header = '\n'.join(bloque.split('\n')[:4])
                    sub_chunks = self._dividir_en_chunks(bloque)
                    for sc in sub_chunks:
                        if not sc.startswith('PERÍODO:'):
                            sc = header + '\n' + sc
                        chunks_texto.append(sc)
                else:
                    chunks_texto.append(bloque)
        else:
            chunks_texto = self._dividir_en_chunks(texto)
        chunks_texto = self._dividir_en_chunks(texto)
        self.stats['chunks_generados'] = len(chunks_texto)
        
        # Crear objetos ProyectoChunk
        chunks = []
        for i, chunk_texto in enumerate(chunks_texto):
            metadata = ProyectoMetadata(
                proyecto_id=proyecto_id,
                titulo=titulo,
                docente_id=docente_id,
                institucion_id=institucion_id,
                grado=str(grado) if grado else "todos",
                tipo=tipo,
                materias=materias,
                año_proyecto=año_proyecto,
                chunk_index=i,
                total_chunks=len(chunks_texto),
            )
            
            chunks.append(ProyectoChunk(
                texto=chunk_texto,
                metadata=metadata,
            ))
        
        return chunks
    
    def get_stats(self) -> dict:
        """Devuelve estadísticas del último procesamiento."""
        return self.stats.copy()


# =============================================================================
# FUNCIÓN DE CONVENIENCIA PARA CARGAR A CHROMADB
# =============================================================================

def preparar_para_chroma(chunks: List[ProyectoChunk]) -> tuple:
    documents = []
    metadatas = []
    ids = []
    
    for chunk in chunks:
        documents.append(chunk.texto)
        
        # Detectar período del chunk si es planificación anual
        periodo = "general"
        if chunk.metadata.tipo == "planificacion_anual":
            texto_lower = chunk.texto.lower()
            if texto_lower.startswith("período:"):
                primera_linea = chunk.texto.split('\n')[0]
                periodo = primera_linea.replace("PERÍODO:", "").strip().lower()
                periodo = periodo.replace(" – ", "-").replace(" - ", "-")
        
        metadatas.append({
            'proyecto_id': chunk.metadata.proyecto_id,
            'titulo': chunk.metadata.titulo,
            'docente_id': chunk.metadata.docente_id,
            'institucion_id': chunk.metadata.institucion_id,
            'grado': chunk.metadata.grado,
            'tipo': chunk.metadata.tipo,
            'materias': ','.join(chunk.metadata.materias),
            'año_proyecto': chunk.metadata.año_proyecto,
            'chunk_index': chunk.metadata.chunk_index,
            'total_chunks': chunk.metadata.total_chunks,
            'periodo': periodo,
        })
        
        ids.append(f"{chunk.metadata.proyecto_id}_chunk{chunk.metadata.chunk_index:03d}")
    
    return documents, metadatas, ids
